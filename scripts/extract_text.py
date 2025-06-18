from pathlib import Path
import sys
import traceback
import time


def debug_log(msg):
    with open("E:/VHC/cloudserver/log.txt", "a", encoding="utf-8") as f:
        f.write(msg + "\n")

debug_log("=== Script Started ===")


def extract_text_from_pdf(file_path):
    import pdfplumber
    text = ""
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
    return text

from docx import Document

def extract_text_from_docx(file_path):
    doc = Document(file_path)
    full_text = []

    # Lấy text từ các đoạn văn
    for para in doc.paragraphs:
        full_text.append(para.text)

    # Lấy text từ bảng
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                row_text.append(cell.text)
            full_text.append('\t'.join(row_text))  # ngăn cách ô bằng tab hoặc dấu khác tùy ý

    return '\n'.join(full_text)


def extract_text_from_txt(file_path):
    with open(file_path, "r", encoding="utf-8") as f:
        return f.read()

def extract_text_from_html(file_path):
    from bs4 import BeautifulSoup
    with open(file_path, "r", encoding="utf-8") as f:
        soup = BeautifulSoup(f, "html.parser")
        return soup.get_text(separator="\n")

def extract_text_from_csv(file_path):
    import csv
    with open(file_path, "r", encoding="utf-8") as f:
        reader = csv.reader(f)
        return "\n".join([", ".join(row) for row in reader])

def extract_text(file_path):
    ext = file_path.suffix.lower()
    if ext == ".pdf":
        return extract_text_from_pdf(file_path)
    elif ext == ".docx":
        return extract_text_from_docx(file_path)
    elif ext in [".txt", ".md"]:
        return extract_text_from_txt(file_path)
    elif ext == ".html":
        return extract_text_from_html(file_path)
    elif ext == ".csv":
        return extract_text_from_csv(file_path)
    else:
        return f"[Unsupported file format: {file_path}]"

if __name__ == "__main__":
    debug_log("=== Script Started ===")
    try:
        if len(sys.argv) < 2:
            print("Usage: python extract_text.py <file_path>")
            sys.exit(1)

        file_path = Path(sys.argv[1])
        debug_log(f"File received: {file_path}")

        if not file_path.exists():
            debug_log(f"[ERROR] File not found: {file_path}")
            print(f"[Error] File not found: {file_path}")
            sys.exit(1)

        text = extract_text(file_path)
        debug_log(f"Extracted {len(text)} characters from {file_path}")

        output_dir = Path("E:/VHC/cloudserver/data/text/extracted")
        output_dir.mkdir(parents=True, exist_ok=True)
        output_file = output_dir / f"{file_path.stem}.txt"

        with open(output_file, "w", encoding="utf-8") as f:
            f.write(text)
        debug_log(f"Saved to: {output_file}")
        print(f"[Saved] Text saved to {output_file}")

    except Exception as e:
        debug_log("[Exception]")
        debug_log(str(e))
        debug_log(traceback.format_exc())
        print(f"[ERROR] {e}")
        sys.exit(1)